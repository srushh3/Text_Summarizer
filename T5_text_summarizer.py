import fitz  
import pytesseract
import re
import urllib.request
import nltk
import requests
from pdf2image import convert_from_path
from transformers import pipeline
from fuzzysearch import find_near_matches
from nltk.tokenize import sent_tokenize
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import RGBColor

nltk.download('punkt')

def pdf_to_images(pdf_path):
    return convert_from_path(pdf_path, dpi=300)

def extract_text_from_image(image):
    return pytesseract.image_to_string(image)

def extract_text_from_pdf(pdf_path):
    pages = pdf_to_images(pdf_path)
    extracted_text = " ".join([extract_text_from_image(page).strip() for page in pages])
    return extracted_text

def extract_text_from_url(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, "html.parser")
    paragraphs = soup.find_all("p")
    text = "\n".join([para.get_text() for para in paragraphs])
    return text.strip()

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="utf-8") as file:
        text = file.read()
    return text

# def preprocess_text(text):
#     """Preprocess the text for summarization."""
#     text = re.sub(r'\[[0-9]*\]', ' ', text)
#     text = re.sub(r'\s+', ' ', text)
#     formatted_text = re.sub('[^a-zA-Z]', ' ', text)
#     return re.sub(r'\s+', ' ', formatted_text)

def summarize_text_t5(text):
    summarizer = pipeline("summarization", model="t5-base", tokenizer="t5-base", framework="tf")
    summary_output = summarizer(text, min_length=5, max_length=950, do_sample=False)
    return summary_output[0]['summary_text']


def fuzzy_match_summary(original_text, summary):
    # Finds approximate matches of summary sentences in the original text using fuzzy matching
    summary_sentences = sent_tokenize(summary)  
    matched_sentences = set()

    for sentence in summary_sentences:
        matches = find_near_matches(sentence, original_text, max_l_dist=20) 
        for match in matches:
            matched_sentences.add(original_text[match.start:match.end])

    return matched_sentences

def save_summary_to_docx(original_text, summary):
    doc = Document()
    doc.add_heading("T5 Text Summarization", level=1)

    # Add Summary Section
    doc.add_paragraph("=== Summary ===")
    doc.add_paragraph(summary)

    # Add Highlighted Text Section
    doc.add_paragraph("\n=== Highlighted Text ===")

    # Use fuzzy matching to find approximate matches of summary sentences in the original text
    matched_sentences = fuzzy_match_summary(original_text, summary)

    # Process each sentence and apply formatting
    para = doc.add_paragraph()
    for sentence in sent_tokenize(original_text):
        run = para.add_run(sentence + " ")  # Add sentence to document
        for matched in matched_sentences:
            if matched in sentence:  # If fuzzy-matched, highlight in red
                run.font.color.rgb = RGBColor(255, 0, 0)
                break  # Stop after finding a match


    # Save the document
    doc.save("T5_summary.docx")
    print("Summary saved in T5_summary.docx")

def main():
    print("Select the input file type:")
    print("1. .txt file")
    print("2. .pdf file")
    print("3. Website URL")
    
    choice = input("Enter your choice (1/2/3): ")
    
    if choice == "1":
        file_path = input("Enter the path of the .txt file: ")
        text = extract_text_from_txt(file_path)
    elif choice == "2":
        file_path = input("Enter the path of the .pdf file: ")
        text = extract_text_from_pdf(file_path)
    elif choice == "3":
        url = input("Enter the website URL: ")
        text = extract_text_from_url(url)
    else:
        print("Invalid choice! Exiting.")
        return
    
    summary = summarize_text_t5(text)
    save_summary_to_docx(text, summary)

if __name__ == "__main__":
    main()