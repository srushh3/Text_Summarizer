import fitz  
import requests
from bs4 import BeautifulSoup
from summarizer import Summarizer
from docx import Document
from docx.shared import RGBColor

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text("text") for page in doc])
    return text

def extract_text_from_url(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, "html.parser")
    paragraphs = soup.find_all("p")
    text = "\n".join([para.get_text() for para in paragraphs])
    return text.strip()

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="utf-8") as file:
        text = file.read()
    return text

def summarize_text_bert(text):
    # Extracts key sentences using BERT-based extractive summarization.
    model = Summarizer()
    summary = model(text, ratio=0.3) 
    return summary

def save_summary_to_docx(original_text, summary):
    doc = Document()
    doc.add_heading("BERT Summary", level=1)
    para = doc.add_paragraph()
    for sentence in original_text.split(". "):
        run = para.add_run(sentence + ". ")
        if sentence in summary:  # Highlight if it is part of the summary
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for summary sentences

    doc.save("Bert_summary.docx")
    print("Summary saved in Bert_summary.docx")

def main():
    print("Select the input file type:")
    print("1. .txt file")
    print("2. .pdf file")
    print("3. Website URL")
    
    choice = input("Enter your choice (1/2/3): ")
    
    if choice == "1":
        file_path = input("Enter the path of the .txt file: ")
        text = extract_text_from_txt(file_path)
    elif choice == "2":
        file_path = input("Enter the path of the .pdf file: ")
        text = extract_text_from_pdf(file_path)
    elif choice == "3":
        url = input("Enter the website URL: ")
        text = extract_text_from_url(url)
    else:
        print("Invalid choice! Exiting.")
        return

    summary = summarize_text_bert(text)
 
    save_summary_to_docx(text, summary)

if __name__ == "__main__":
    main()
